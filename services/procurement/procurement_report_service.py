"""
구매견적서 생성 서비스 — DOCX / PDF
"""
import io
from datetime import datetime


# ── 공통 데이터 정규화 ────────────────────────────────────────

def _normalize(data: dict) -> dict:
    """프론트에서 넘어온 데이터를 안전하게 정규화"""
    candidates = data.get("top_candidates") or []
    selected_rank = int(data.get("selected_candidate_rank") or 1)

    return {
        "order_id":      data.get("order_id", "-"),
        "department":    data.get("department", ""),
        "item_name":     data.get("item_name", ""),
        "quantity":      int(data.get("quantity") or 1),
        "unit_price":    int(data.get("unit_price") or 0),
        "total_amount":  int(data.get("total_amount") or 0),
        "vendor":        data.get("vendor", ""),
        "account_code":  data.get("account_code", ""),
        "status":        data.get("status", "승인대기"),
        "created_at":    str(data.get("created_at") or "")[:10] or datetime.now().strftime("%Y-%m-%d"),
        "report_text":   data.get("report_text", ""),
        "candidates":    candidates,
        "selected_rank": selected_rank,
    }


# ── DOCX 생성 ────────────────────────────────────────────────

def generate_estimate_docx(data: dict) -> bytes:
    """구매견적서 DOCX 생성"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = _normalize(data)
    doc = Document()

    # ── 여백 설정 ─────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 제목 ──────────────────────────────────────────────────
    title = doc.add_heading("구  매  견  적  서", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(20)

    doc.add_paragraph()

    # ── 기본 정보 테이블 ──────────────────────────────────────
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"

    cells = [
        ("작 성 일", d["created_at"]),
        ("요청 부서", d["department"]),
        ("주문 번호", f"#{d['order_id']}"),
        ("처 리 상 태", d["status"]),
    ]
    for i, (label, val) in enumerate(cells):
        row_idx, col_idx = divmod(i, 2)
        info.rows[row_idx].cells[col_idx * 2].text     = label
        info.rows[row_idx].cells[col_idx * 2 + 1].text = val

    _style_header_cells(info, col_indices=[0, 2])
    doc.add_paragraph()

    # ── 구매 품목 정보 ────────────────────────────────────────
    doc.add_heading("구매 품목 정보", level=2)
    item_tbl = doc.add_table(rows=2, cols=6)
    item_tbl.style = "Table Grid"

    headers = ["품목명", "수량", "단가", "총액", "계정과목", "공급업체"]
    values  = [
        d["item_name"],
        str(d["quantity"]),
        f"{d['unit_price']:,}원",
        f"{d['total_amount']:,}원",
        d["account_code"],
        d["vendor"],
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        item_tbl.rows[0].cells[i].text = h
        item_tbl.rows[1].cells[i].text = v

    _style_header_row(item_tbl, row_idx=0)
    doc.add_paragraph()

    # ── 상위 3개 추천 상품 ────────────────────────────────────
    if d["candidates"]:
        doc.add_heading("상위 3개 추천 상품 비교", level=2)
        cand_tbl = doc.add_table(rows=1 + len(d["candidates"]), cols=5)
        cand_tbl.style = "Table Grid"

        col_headers = ["순위", "상품명", "단가", "판매처", "추천 이유"]
        for i, h in enumerate(col_headers):
            cand_tbl.rows[0].cells[i].text = h

        _style_header_row(cand_tbl, row_idx=0)

        for ri, c in enumerate(d["candidates"], start=1):
            rank   = c.get("rank", ri)
            name   = c.get("name", "")
            price  = c.get("price", 0)
            vendor = c.get("vendor", "")
            reason = c.get("reason", "")
            url    = c.get("url", "")

            row = cand_tbl.rows[ri]
            row.cells[0].text = f"{'★' if rank == d['selected_rank'] else '  '} {rank}위"
            row.cells[1].text = name
            row.cells[2].text = f"{price:,}원" if price else "-"
            row.cells[3].text = vendor
            row.cells[4].text = reason

            # 구매 링크 하이퍼링크를 상품명 셀에 추가
            if url:
                _add_hyperlink(row.cells[1].paragraphs[0], url, "구매 링크")

            # 최종 선택 행 배경 강조
            if rank == d["selected_rank"]:
                _shade_row(row, "FFF3E0")

        doc.add_paragraph()

        # 최종 선택 상품 구매 링크 별도 표기
        selected = next((c for c in d["candidates"] if c.get("rank") == d["selected_rank"]), None)
        if selected and selected.get("url"):
            p = doc.add_paragraph("최종 선택 구매 링크: ")
            _add_hyperlink(p, selected["url"], selected.get("vendor", "구매 링크"))

        doc.add_paragraph()

    # ── AI 구매 분석 보고서 ───────────────────────────────────
    if d["report_text"]:
        doc.add_heading("AI 구매 분석 보고서", level=2)
        for line in d["report_text"].splitlines():
            doc.add_paragraph(line if line.strip() else "")

    # ── 바이트 반환 ───────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF 생성 ──────────────────────────────────────────────────

def generate_estimate_pdf(data: dict) -> bytes:
    """구매견적서 PDF 생성 (ReportLab + 한국어 CID 폰트)"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, HRFlowable,
    )
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase import pdfmetrics

    # 한국어 폰트 등록
    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
    pdfmetrics.registerFont(UnicodeCIDFont("HYGothic-Medium"))
    KO_SERIF = "HYSMyeongJo-Medium"
    KO_SANS  = "HYGothic-Medium"

    d = _normalize(data)
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=3.0 * cm,
        rightMargin=2.5 * cm,
    )

    # ── 스타일 정의 ───────────────────────────────────────────
    ORANGE = colors.HexColor("#EA580C")
    LIGHT_ORANGE = colors.HexColor("#FFF7ED")
    GRAY_BG = colors.HexColor("#F9FAFB")

    s_title = ParagraphStyle("title",   fontName=KO_SERIF, fontSize=20, leading=28, alignment=1, spaceAfter=6)
    s_h2    = ParagraphStyle("h2",      fontName=KO_SANS,  fontSize=13, leading=18, textColor=ORANGE, spaceBefore=14, spaceAfter=6)
    s_body  = ParagraphStyle("body",    fontName=KO_SANS,  fontSize=10, leading=16, spaceAfter=4)
    s_small = ParagraphStyle("small",   fontName=KO_SANS,  fontSize=9,  leading=14, textColor=colors.gray)
    s_link  = ParagraphStyle("link",    fontName=KO_SANS,  fontSize=9,  leading=14, textColor=colors.blue)

    story = []

    # ── 제목 ──────────────────────────────────────────────────
    story.append(Paragraph("구  매  견  적  서", s_title))
    story.append(HRFlowable(width="100%", thickness=2, color=ORANGE, spaceAfter=10))

    # ── 기본 정보 테이블 ──────────────────────────────────────
    info_data = [
        ["작 성 일", d["created_at"],    "요청 부서", d["department"]],
        ["주문 번호", f"#{d['order_id']}", "처리 상태", d["status"]],
    ]
    info_tbl = Table(info_data, colWidths=[3 * cm, 5.5 * cm, 3 * cm, 5.5 * cm])
    info_tbl.setStyle(TableStyle([
        ("FONTNAME",    (0, 0), (-1, -1), KO_SANS),
        ("FONTSIZE",    (0, 0), (-1, -1), 10),
        ("BACKGROUND",  (0, 0), (0, -1), GRAY_BG),
        ("BACKGROUND",  (2, 0), (2, -1), GRAY_BG),
        ("FONTNAME",    (0, 0), (0, -1), KO_SANS),
        ("FONTNAME",    (2, 0), (2, -1), KO_SANS),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUND", (0, 0), (-1, -1), [colors.white, colors.white]),
        ("TOPPADDING",  (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 12))

    # ── 구매 품목 ─────────────────────────────────────────────
    story.append(Paragraph("구매 품목 정보", s_h2))
    item_data = [
        ["품목명", "수량", "단가", "총액", "계정과목", "공급업체"],
        [
            d["item_name"],
            str(d["quantity"]),
            f"{d['unit_price']:,}원",
            f"{d['total_amount']:,}원",
            d["account_code"],
            d["vendor"],
        ],
    ]
    item_tbl = Table(item_data, colWidths=[4 * cm, 1.5 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm, 3.5 * cm])
    item_tbl.setStyle(TableStyle([
        ("FONTNAME",      (0, 0), (-1, -1), KO_SANS),
        ("FONTSIZE",      (0, 0), (-1, -1), 9),
        ("BACKGROUND",    (0, 0), (-1, 0),  ORANGE),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",      (0, 0), (-1, 0),  KO_SANS),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
        ("BACKGROUND",    (0, 1), (-1, 1),  LIGHT_ORANGE),
    ]))
    story.append(item_tbl)
    story.append(Spacer(1, 12))

    # ── 상위 3개 추천 상품 ────────────────────────────────────
    if d["candidates"]:
        story.append(Paragraph("상위 3개 추천 상품 비교", s_h2))

        cand_rows = [["순위", "상품명", "단가", "판매처", "추천 이유", "구매 링크"]]
        for c in d["candidates"]:
            rank   = c.get("rank", 1)
            mark   = "★ " if rank == d["selected_rank"] else f"{rank}위"
            price  = c.get("price", 0)
            url    = c.get("url", "")
            link_p = Paragraph(
                f'<link href="{url}" color="blue">링크</link>' if url else "-",
                s_link,
            )
            cand_rows.append([
                f"{mark}{rank}위" if mark != f"{rank}위" else mark,
                c.get("name", ""),
                f"{price:,}원" if price else "-",
                c.get("vendor", ""),
                c.get("reason", ""),
                link_p,
            ])

        cand_tbl = Table(cand_rows, colWidths=[1.5 * cm, 5 * cm, 2.5 * cm, 3 * cm, 3.5 * cm, 1.5 * cm])
        style_cmds = [
            ("FONTNAME",      (0, 0), (-1, -1), KO_SANS),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("BACKGROUND",    (0, 0), (-1, 0),  ORANGE),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("GRID",          (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("TOPPADDING",    (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("WORDWRAP",      (0, 0), (-1, -1), True),
        ]
        # 최종 선택 행 배경
        for ri, c in enumerate(d["candidates"], start=1):
            if c.get("rank") == d["selected_rank"]:
                style_cmds.append(("BACKGROUND", (0, ri), (-1, ri), LIGHT_ORANGE))

        cand_tbl.setStyle(TableStyle(style_cmds))
        story.append(cand_tbl)
        story.append(Spacer(1, 6))

        # 최종 선택 안내
        selected = next((c for c in d["candidates"] if c.get("rank") == d["selected_rank"]), None)
        if selected:
            url = selected.get("url", "")
            link_text = f' (<link href="{url}" color="blue">구매 링크</link>)' if url else ""
            story.append(Paragraph(
                f"✔ 최종 선택: {selected.get('name', '')} — {selected.get('vendor', '')}{link_text}",
                s_body,
            ))
        story.append(Spacer(1, 12))

    # ── AI 구매 분석 보고서 ───────────────────────────────────
    if d["report_text"]:
        story.append(Paragraph("AI 구매 분석 보고서", s_h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceAfter=6))
        for line in d["report_text"].splitlines():
            text = line.strip()
            story.append(Paragraph(text if text else " ", s_body))

    doc.build(story)
    return buf.getvalue()


# ── DOCX 헬퍼 함수들 ─────────────────────────────────────────

def _style_header_cells(table, col_indices: list):
    """지정 열 셀 배경색 회색 처리"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for row in table.rows:
        for ci in col_indices:
            cell = row.cells[ci]
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F3F4F6")
            tc_pr.append(shd)


def _style_header_row(table, row_idx: int = 0):
    """지정 행을 주황색 헤더로 스타일"""
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for cell in table.rows[row_idx].cells:
        # 배경색
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "EA580C")
        tc_pr.append(shd)
        # 글자색 흰색
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if not para.runs:
                run = para.add_run(para.text)
                para.clear()
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                para.add_run(run.text)


def _shade_row(row, hex_color: str):
    """행 배경색 설정"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tc_pr.append(shd)


def _add_hyperlink(paragraph, url: str, text: str):
    """DOCX 하이퍼링크 추가"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
