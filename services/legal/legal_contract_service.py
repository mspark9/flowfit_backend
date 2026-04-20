"""
법무 계약 서비스 — 계약서 AI 검토 / 계약서 초안 생성 / DOCX 다운로드
legal_document_chunks RAG를 활용하여 사내 법령·사규 기준으로 분석·생성
"""
import io
import json
import re

from openai import OpenAI

from config import settings
from database import get_connection
from services.common.document_parser import (
    encode_image_base64,
    extract_document_text,
    is_image_file,
)
from services.common.rag_utils import embed_text, select_top_chunks_by_vector

client = OpenAI(api_key=settings.openai_api_key)


# ── 내부 헬퍼 ─────────────────────────────────────────────────

def _load_rag_context(query: str, top_k: int = 5) -> tuple[str, list[str]]:
    """
    질의 임베딩 → legal_document_chunks 코사인 유사도 검색.
    반환: (context_text, source_file_names)
    법무 문서가 없을 경우 빈 문자열 반환 (기능 자체는 계속 동작)
    """
    conn = get_connection()
    cur = conn.cursor()

    try:
        # 테이블 존재 여부 확인
        cur.execute(
            """
            SELECT EXISTS (
                SELECT 1 FROM information_schema.tables
                WHERE table_name = 'legal_document_chunks'
            )
            """
        )
        if not cur.fetchone()[0]:
            return "", []

        cur.execute(
            """
            SELECT c.file_name, c.chunk_text, c.embedding
            FROM legal_document_chunks c
            JOIN legal_documents d ON d.id = c.document_id
            WHERE d.deleted_at IS NULL AND d.is_active = TRUE
            ORDER BY c.document_id, c.chunk_index
            """
        )
        rows = cur.fetchall()
    finally:
        cur.close()
        conn.close()

    if not rows:
        return "", []

    # 질의 임베딩 → 유사 청크 선택
    query_embedding = embed_text(query)
    chunks = [
        {"file_name": row[0], "chunk_text": row[1], "embedding": list(row[2])}
        for row in rows
    ]
    top_chunks = select_top_chunks_by_vector(query_embedding, chunks, top_k=top_k)

    context = "\n\n".join(
        f"[참조: {item['file_name']}]\n{item['chunk_text']}"
        for item in top_chunks
    )
    sources = list(dict.fromkeys(item["file_name"] for item in top_chunks))
    return context, sources


def _rag_context_section(context: str, sources: list[str]) -> str:
    """RAG 컨텍스트가 있을 때 프롬프트에 삽입할 섹션 생성"""
    if not context:
        return ""
    return (
        f"\n\n[사내 법령·사규 참조 문서: {', '.join(sources)}]\n"
        f"{context}\n"
        "위 참조 문서가 있을 경우 해당 기준을 우선 적용하여 분석하세요.\n"
    )


# ── 계약서 검토 공통 프롬프트 ────────────────────────────────

_REVIEW_SYSTEM_PROMPT = (
    "당신은 사내 법무·컴플라이언스 전문가입니다. "
    "계약서를 분석하여 리스크 조항을 탐지하고 수정안을 제안하세요. "
    "risk_level은 danger(즉각 수정 필요), warning(검토 권고), safe(정상) 세 가지입니다. "
    "응답은 JSON만 반환하세요."
)

_REVIEW_FORMAT = (
    "\n[반환 형식]\n"
    "{\n"
    '  "overall_risk": "danger|warning|safe",\n'
    '  "summary": "전체 리스크 요약 (2~3문장)",\n'
    '  "clauses": [\n'
    "    {\n"
    '      "id": 1,\n'
    '      "risk_level": "danger|warning|safe",\n'
    '      "title": "조항 제목",\n'
    '      "article": "제 X조 제 Y항 (없으면 빈 문자열)",\n'
    '      "original_text": "원문 발췌",\n'
    '      "ai_comment": "AI 분석 및 리스크 설명",\n'
    '      "suggestion": "수정 제안 (safe이면 null)"\n'
    "    }\n"
    "  ]\n"
    "}"
)


def _parse_review_response(response) -> dict:
    """GPT 응답 파싱 및 조항 id 보정"""
    payload = json.loads(response.choices[0].message.content)
    clauses = payload.get("clauses", [])
    for i, clause in enumerate(clauses):
        clause["id"] = i + 1
        if clause.get("suggestion") == "":
            clause["suggestion"] = None
    return {
        "overall_risk": payload.get("overall_risk", "safe"),
        "summary": payload.get("summary", ""),
        "clauses": clauses,
    }


# ── 계약서 검토 ───────────────────────────────────────────────

def _review_text_contract(filename: str, file_bytes: bytes) -> dict:
    """텍스트 기반 계약서 검토 (pdf, docx, hwp, txt)"""
    contract_text = extract_document_text(filename, file_bytes)
    rag_context, rag_sources = _load_rag_context(contract_text[:1000], top_k=5)
    rag_section = _rag_context_section(rag_context, rag_sources)

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": _REVIEW_SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"[계약서 내용]\n{contract_text[:6000]}\n"
                    f"{rag_section}"
                    f"{_REVIEW_FORMAT}"
                ),
            },
        ],
        max_tokens=2000,
        temperature=0.2,
    )

    result = _parse_review_response(response)
    result["rag_sources"] = rag_sources
    return result


def _review_image_contract(filename: str, file_bytes: bytes) -> dict:
    """이미지 계약서 검토 — Vision API로 이미지 직접 분석"""
    b64_image, mime_type = encode_image_base64(filename, file_bytes)

    # 이미지 자체에서 내용 파악이 불가하므로 범용 쿼리로 RAG 검색
    rag_context, rag_sources = _load_rag_context("계약서 리스크 검토 손해배상 해지 조항", top_k=5)
    rag_section = _rag_context_section(rag_context, rag_sources)

    # Vision API: 텍스트 + 이미지를 content 배열로 전달
    user_content = [
        {
            "type": "text",
            "text": (
                "첨부된 계약서 이미지를 분석하여 리스크 조항을 탐지하세요.\n"
                "이미지에서 텍스트를 읽어 조항별로 분석해 주세요.\n"
                f"{rag_section}"
                f"{_REVIEW_FORMAT}"
            ),
        },
        {
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime_type};base64,{b64_image}",
                "detail": "high",   # 고해상도 분석 모드
            },
        },
    ]

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": _REVIEW_SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        max_tokens=2000,
        temperature=0.2,
    )

    result = _parse_review_response(response)
    result["rag_sources"] = rag_sources
    return result


def review_contract(filename: str, file_bytes: bytes) -> dict:
    """
    계약서 AI 리스크 검토 — 파일 형식 자동 감지:
    - 이미지 (jpg/png/webp/gif): Vision API로 직접 분석
    - 문서 (pdf/docx/hwp/txt): 텍스트 추출 후 분석
    두 경우 모두 RAG로 사내 법령·사규 참조
    """
    if not file_bytes:
        raise ValueError("파일이 비어 있습니다.")

    if is_image_file(filename):
        return _review_image_contract(filename, file_bytes)
    return _review_text_contract(filename, file_bytes)


# ── 계약서 초안 생성 ──────────────────────────────────────────

CONTRACT_TYPE_LABELS = {
    "nda":         "NDA (비밀유지계약)",
    "service":     "용역계약서",
    "employment":  "근로계약서",
    "partnership": "업무협약서 (MOU)",
    "purchase":    "물품 구매계약서",
}


def generate_contract_draft(form: dict) -> dict:
    """
    계약서 초안 생성:
    1. 계약 목적·유형으로 RAG에서 관련 법령·사규 검색
    2. GPT가 해당 법적 기준을 반영한 계약서 초안 생성
    """
    contract_type = form.get("contract_type", "")
    party_a = form.get("party_a", "")
    party_b = form.get("party_b", "")
    purpose = form.get("purpose", "")
    amount = form.get("amount", "")
    start_date = form.get("start_date", "")
    end_date = form.get("end_date", "")
    extra = form.get("extra", "")

    if not (contract_type and party_a and party_b and purpose):
        raise ValueError("계약 유형, 갑, 을, 계약 목적은 필수 입력 항목입니다.")

    type_label = CONTRACT_TYPE_LABELS.get(contract_type, contract_type)

    # RAG: 계약 유형 + 목적 기반으로 관련 법령·사규 검색
    rag_query = f"{type_label} {purpose}"
    rag_context, rag_sources = _load_rag_context(rag_query, top_k=5)
    rag_section = _rag_context_section(rag_context, rag_sources)

    # 입력 조건 정리
    conditions = (
        f"계약 유형: {type_label}\n"
        f"갑(발주/위탁사): {party_a}\n"
        f"을(수주/수탁사): {party_b}\n"
        f"계약 목적/업무 범위: {purpose}\n"
    )
    if amount:
        conditions += f"계약 금액: {amount}원\n"
    if start_date:
        conditions += f"계약 시작일: {start_date}\n"
    if end_date:
        conditions += f"계약 종료일: {end_date}\n"
    if extra:
        conditions += f"특이 사항: {extra}\n"

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": (
                    "당신은 사내 법무·컴플라이언스 전문가입니다. "
                    "입력된 계약 조건을 바탕으로 대한민국 법률에 부합하는 계약서 초안을 한국어로 작성하세요. "
                    "사내 법령·사규 참조 문서가 제공된 경우 해당 기준을 우선 반영하세요. "
                    "조항은 제1조부터 순서대로 작성하고 마지막에 서명란을 포함하세요. "
                    "응답은 JSON만 반환하세요."
                ),
            },
            {
                "role": "user",
                "content": (
                    f"[계약 조건]\n{conditions}"
                    f"{rag_section}"
                    "\n[반환 형식]\n"
                    "{\n"
                    '  "draft": "완성된 계약서 전문 (줄바꿈 포함 plain text)",\n'
                    '  "note": "생성 시 참고한 사항 또는 주의사항 (1~2문장)"\n'
                    "}"
                ),
            },
        ],
        max_tokens=2500,
        temperature=0.3,
    )

    payload = json.loads(response.choices[0].message.content)

    return {
        "draft": payload.get("draft", ""),
        "note": payload.get("note", ""),
        "rag_sources": rag_sources,
    }


# ── DOCX 다운로드 ─────────────────────────────────────────────

# 조항 헤더 패턴 (제1조, 제 2조 (목적) 등)
_ARTICLE_RE = re.compile(r"^제\s*\d+\s*조")

def draft_to_docx(draft_text: str, contract_title: str = "계약서") -> bytes:
    """
    계약서 plain text → DOCX 바이트 변환.
    - 첫 번째 비어 있지 않은 줄: 제목 (Heading 1)
    - '제X조'로 시작하는 줄: 조항 제목 (Heading 2)
    - 나머지: 본문 단락 (Normal)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 문서 기본 여백 설정 (좌우 2.5cm)
    for section in doc.sections:
        section.left_margin = section.right_margin = int(2.54 * 914400 / 2.54)  # 2.54cm

    lines = draft_text.splitlines()
    first_line_done = False

    for raw_line in lines:
        line = raw_line.strip()

        # 첫 번째 비어 있지 않은 줄 → 계약서 제목
        if not first_line_done and line:
            title_para = doc.add_heading(line, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_line_done = True
            continue

        # 빈 줄 → 빈 단락 (간격)
        if not line:
            doc.add_paragraph("")
            continue

        # 조항 헤더 (제X조...)
        if _ARTICLE_RE.match(line):
            doc.add_heading(line, level=2)
            continue

        # 일반 본문
        doc.add_paragraph(line)

    # 메모리 버퍼에 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
